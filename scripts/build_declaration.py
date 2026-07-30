"""Generate the Declaration of Interest file for upload.

Elsevier's submission form offers a checkbox for "no competing interests", but
several journals require the declaration as an uploaded document regardless. Its
own tool produces a Word file; this produces the same document, in Word and PDF,
using Elsevier's standard wording and tick-box layout so it is recognisable to
the editorial office.

The declaration itself is a single author with nothing to declare, consistent
with the Declaration of competing interest section inside the manuscript.

Writes paper/declaration_of_interest.docx and .pdf

Run:  python -X utf8 scripts/build_declaration.py
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt
from weasyprint import HTML

ROOT = Path(__file__).resolve().parents[1]
PAPER = ROOT / "paper"

TITLE = ("A neural surrogate for the natural-gas compressibility factor: "
         "reference-equation accuracy from synthetic training data alone, "
         "benchmarked on 2426 independent laboratory measurements")
AUTHOR = "Emir Borovac"

HEADING = "Declaration of interests"

# Elsevier's standard two-option form. The selected box is marked with a filled
# ballot box, exactly as their own tool renders it.
OPT_NONE = ("☒ The authors declare that they have no known competing "
            "financial interests or personal relationships that could have "
            "appeared to influence the work reported in this paper.")
OPT_SOME = ("☐ The authors declare the following financial "
            "interests/personal relationships which may be considered as "
            "potential competing interests:")


def build_docx(path: Path) -> None:
    d = Document()
    st = d.styles["Normal"]
    st.font.name = "Calibri"
    st.font.size = Pt(11)

    h = d.add_paragraph()
    r = h.add_run(HEADING)
    r.bold = True
    r.font.size = Pt(14)

    d.add_paragraph()
    p = d.add_paragraph()
    p.add_run("Manuscript title: ").bold = True
    p.add_run(TITLE)

    p = d.add_paragraph()
    p.add_run("Author: ").bold = True
    p.add_run(f"{AUTHOR} (sole author)")

    d.add_paragraph()
    d.add_paragraph(OPT_NONE)
    d.add_paragraph()
    d.add_paragraph(OPT_SOME)
    d.add_paragraph()
    d.add_paragraph("(none)")
    d.add_paragraph()
    d.add_paragraph(
        "The author further declares that he has not served in an editorial "
        "capacity for the journal to which this manuscript is submitted, and "
        "that this work received no specific funding from any agency in the "
        "public, commercial or not-for-profit sectors.")
    d.add_paragraph()
    p = d.add_paragraph()
    p.add_run("Signed: ").bold = True
    p.add_run(AUTHOR)
    d.add_paragraph("Independent Researcher")
    d.add_paragraph("emirborovac@rippleit-co.com")

    d.save(str(path))


def build_pdf(path: Path) -> None:
    html = f"""<!doctype html><html><head><meta charset="utf-8"><style>
    @page {{ size: A4; margin: 25mm; }}
    body {{ font: 11pt/1.6 Calibri, Arial, sans-serif; color:#111; }}
    h1 {{ font-size: 14pt; margin: 0 0 18pt; }}
    p {{ margin: 0 0 10pt; }}
    .lbl {{ font-weight: bold; }}
    .opt {{ margin: 14pt 0; }}
    .sig {{ margin-top: 26pt; }}
    </style></head><body>
    <h1>{HEADING}</h1>
    <p><span class="lbl">Manuscript title:</span> {TITLE}</p>
    <p><span class="lbl">Author:</span> {AUTHOR} (sole author)</p>
    <p class="opt">{OPT_NONE}</p>
    <p class="opt">{OPT_SOME}</p>
    <p>(none)</p>
    <p>The author further declares that he has not served in an editorial
    capacity for the journal to which this manuscript is submitted, and that
    this work received no specific funding from any agency in the public,
    commercial or not-for-profit sectors.</p>
    <p class="sig"><span class="lbl">Signed:</span> {AUTHOR}<br>
    Independent Researcher<br>emirborovac@rippleit-co.com</p>
    </body></html>"""
    HTML(string=html).write_pdf(str(path))


def main() -> None:
    dx = PAPER / "declaration_of_interest.docx"
    pf = PAPER / "declaration_of_interest.pdf"
    build_docx(dx)
    build_pdf(pf)
    for f in (dx, pf):
        print(f"  {f.stat().st_size:>7,} bytes  {f.relative_to(ROOT)}")

    # the uploaded declaration must agree with the manuscript's own statement
    tex = (PAPER / "manuscript.tex").read_text(encoding="utf-8")
    if "no known competing financial interests" not in tex:
        raise SystemExit("manuscript's competing-interest section does not "
                         "match the declaration wording")
    print("  consistent with the manuscript's Declaration of competing interest")


if __name__ == "__main__":
    main()
